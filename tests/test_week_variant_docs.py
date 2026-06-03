from __future__ import annotations

import sys
import tempfile
import unittest
import subprocess
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
PYTHON = Path("/Users/a123/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3")
sys.path.insert(0, str(ROOT / "scripts"))


class WeekVariantDocTests(unittest.TestCase):
    def test_generate_three_word_variant_docs_from_raw_material(self) -> None:
        import generate_week_variant_docs

        source = ROOT / "weeks/2026-06-01_2026-06-05/inputs/word/raw/金葵花债市周度复盘20260528.docx"
        with tempfile.TemporaryDirectory() as tmp:
            output_dir = Path(tmp)
            outputs = generate_week_variant_docs.build_variant_docs(source, output_dir)

            self.assertEqual(set(outputs), {"原版", "固收+", "债市"})
            for variant, path in outputs.items():
                with self.subTest(variant=variant):
                    self.assertTrue(path.exists())
                    paragraphs = [p.text.strip() for p in Document(path).paragraphs if p.text.strip()]
                    first_section = paragraphs.index("第一板块：（最顶部）投资建议：")
                    top_intro = paragraphs[first_section + 1]
                    self.assertIn("5月22日至5月28日", "\n".join(paragraphs))
                    self.assertIn("数据来源：wind，截至2026年5月28日", paragraphs)
                    self.assertIn("本周资金利率：上行", paragraphs)
                    self.assertEqual(len(Document(path).tables), 2)

                    if variant == "原版":
                        self.assertIn("本周债券收益率全期限下行2-4BP", top_intro)
                        self.assertIn("关注长端超额下行空间", paragraphs[first_section + 2])
                    elif variant == "固收+":
                        self.assertIn("固收+策略的配置价值", top_intro)
                        self.assertIn("股债偏强共振", top_intro)
                    else:
                        self.assertIn("继续持有纯债产品", top_intro)
                        self.assertIn("债市多头趋势或仍能延续", top_intro)

    def test_extract_current_week_full_word_variant_for_long_image_fields(self) -> None:
        import extract_word_psd_content

        rules = extract_word_psd_content.load_rules()
        docx_path = ROOT / "weeks/2026-06-01_2026-06-05/inputs/word/variants/债市周观察原版.docx"

        data = extract_word_psd_content.build_variant("原版", docx_path, rules)

        self.assertEqual(data["meta"]["date"], "2026-5-28")
        self.assertEqual(data["meta"]["week_range"], "5月22日至5月28日")
        self.assertEqual(data["fields"]["subtitle"], "收益率全期限下行")
        self.assertEqual(data["fields"]["market_sentence"], "5月22日至5月28日，国债收益率全期限下行。")
        self.assertIn("上证指数窄幅震荡于4100点一线", data["fields"]["risk_preference"])
        self.assertIn("DR007收1.3578%", data["fields"]["funding"])
        self.assertIn("1年期国股行存单发行利率下行至1.44%", data["fields"]["funding"])
        self.assertEqual(data["assets"]["yield_chart_xml"], "word/charts/chart1.xml")
        self.assertEqual(data["assets"]["fund_chart_xml"], "word/charts/chart7.xml")

    def test_current_week_asset_builder_selects_fund_rate_chart(self) -> None:
        import prepare_basic_v0_assets

        self.assertEqual(
            prepare_basic_v0_assets.chart_name_for_series(("回购加权利率:1天", "回购加权利率:7天")),
            "chart7.xml",
        )
        self.assertEqual(
            prepare_basic_v0_assets.chart_name_for_series(("中债国债到期收益率:1年",)),
            "chart1.xml",
        )
        self.assertEqual(str(prepare_basic_v0_assets.source_date()), "2026-05-28")

    def test_native_v3_uses_current_week_risk_text_and_output_date_slug(self) -> None:
        subprocess.run(
            [str(PYTHON), "scripts/prepare_native_v3_content.py", "--variant", "原版"],
            cwd=ROOT,
            check=True,
        )

        data = __import__("json").loads(
            (ROOT / "weeks/2026-06-01_2026-06-05/work/native_v3/content.json").read_text(
                encoding="utf-8"
            )
        )
        risk_text = data["native_text_blocks"]["risk_preference"]["text"]
        self.assertIn("科技概念短线回调", risk_text)
        self.assertIn("消费板块低位反弹", risk_text)
        self.assertNotIn("美伊停战协议", risk_text)
        self.assertIn("20260528_原版_原生文本v3", data["output_png"])
        self.assertFalse(data["native_text_blocks"]["market_sentence"]["overflow"])


if __name__ == "__main__":
    unittest.main()
