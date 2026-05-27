from __future__ import annotations

import json
import re
import zipfile
import xml.etree.ElementTree as ET
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont

from project_paths import EDITED_DOCX, OUTPUT_DIR, PSD_TEMPLATE, REFERENCE_LONG_IMAGE_DIR, WORK_DIR as WEEK_WORK_DIR

DOCX = EDITED_DOCX
PSD = PSD_TEMPLATE
OLD_ORIGINAL = REFERENCE_LONG_IMAGE_DIR / "原版.png"
WORK_DIR = WEEK_WORK_DIR / "basic_v0"
ASSET_DIR = WORK_DIR / "assets"
OUT_DIR = OUTPUT_DIR / "basic_v0"

FONT_HEITI = "/System/Library/Fonts/STHeiti Medium.ttc"
FONT_SONG = "/System/Library/Fonts/Supplemental/Songti.ttc"

X0 = 3042
LONG_W = 1125
MARGIN = 40
CARD_W = 1045

BLACK = (26, 24, 22)
RED = (232, 31, 31)
GREEN = (0, 172, 96)
GREY = (96, 96, 96)
GRID = (218, 218, 218)
WHITE = (255, 255, 255)
CREAM = (255, 250, 241)
DARK = (33, 25, 22)
GOLD = (216, 160, 78)
LIGHT_GOLD = (248, 217, 172)


def font(size: int, song: bool = False) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(FONT_SONG if song else FONT_HEITI, size)


def text_width(draw: ImageDraw.ImageDraw, text: str, ft: ImageFont.FreeTypeFont) -> float:
    return draw.textlength(text, font=ft)


def wrap_rich_text(segments, size: int, max_width: int, line_height: int):
    measure = ImageDraw.Draw(Image.new("RGB", (10, 10)))
    ft = font(size)
    lines: list[list[tuple[str, tuple[int, int, int]]]] = []
    line: list[tuple[str, tuple[int, int, int]]] = []
    width = 0.0
    for text, color in segments:
        for ch in text:
            if ch == "\n":
                lines.append(line)
                line = []
                width = 0.0
                continue
            w = text_width(measure, ch, ft)
            if line and width + w > max_width:
                lines.append(line)
                line = []
                width = 0.0
            line.append((ch, color))
            width += w
    if line:
        lines.append(line)
    return lines, len(lines) * line_height


def rich_text_layers(name: str, x: int, y: int, segments, size: int, max_width: int, line_height: int):
    measure = ImageDraw.Draw(Image.new("RGB", (10, 10)))
    ft = font(size)
    lines, height = wrap_rich_text(segments, size, max_width, line_height)
    layers = []
    for line_index, line in enumerate(lines):
        cx = x
        cy = y + line_index * line_height
        run_text = ""
        run_color = None
        run_x = cx

        def flush():
            nonlocal run_text, run_color, run_x
            if run_text:
                layers.append(
                    {
                        "type": "text",
                        "name": f"{name}_{line_index + 1:02d}",
                        "text": run_text,
                        "x": round(run_x),
                        "y": round(cy),
                        "size": size,
                        "color": run_color,
                    }
                )
            run_text = ""
            run_color = None

        for ch, color in line:
            if run_color is None:
                run_color = color
                run_x = cx
            if color != run_color:
                flush()
                run_color = color
                run_x = cx
            run_text += ch
            cx += text_width(measure, ch, ft)
        flush()
    return layers, height


def plain_text_layers(name: str, x: int, y: int, text: str, size: int, max_width: int, line_height: int, color=BLACK):
    return rich_text_layers(name, x, y, [(text, color)], size, max_width, line_height)


def get_paragraphs():
    doc = Document(DOCX)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()], doc


def parse_table(doc: Document):
    return [[cell.text.strip() for cell in row.cells] for row in doc.tables[0].rows]


def excel_date(n: str) -> datetime:
    return datetime(1899, 12, 30) + timedelta(days=float(n))


def parse_chart(chart_name: str):
    ns = {"c": "http://schemas.openxmlformats.org/drawingml/2006/chart"}
    with zipfile.ZipFile(DOCX) as z:
        root = ET.fromstring(z.read(f"word/charts/{chart_name}"))
    series = []
    for ser in root.findall(".//c:ser", ns):
        tx = ser.find(".//c:tx//c:v", ns)
        name = tx.text if tx is not None else "series"
        cats = [pt.find("c:v", ns).text for pt in ser.findall(".//c:cat//c:pt", ns)]
        vals = [float(pt.find("c:v", ns).text) for pt in ser.findall(".//c:val//c:pt", ns)]
        pairs = [(excel_date(c), v) for c, v in zip(cats, vals)]
        pairs = [(d, v) for d, v in pairs if d.date() <= datetime(2026, 5, 21).date()]
        pairs.sort(key=lambda x: x[0])
        series.append((name, pairs))
    return series


def rounded_asset(path: Path, size: tuple[int, int], radius: int, fill, outline=None, width: int = 2):
    im = Image.new("RGBA", size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(im)
    box = (0, 0, size[0] - 1, size[1] - 1)
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)
    im.save(path)


def frame_asset(path: Path, width: int, height: int, title_width: int = 430):
    im = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    draw = ImageDraw.Draw(im)
    draw.rounded_rectangle((0, 70, width - 1, height - 1), radius=28, fill=CREAM, outline=(245, 228, 196), width=2)
    draw.rounded_rectangle((0, 0, title_width, 110), radius=32, fill=GOLD)
    draw.rectangle((0, 60, title_width - 50, 125), fill=GOLD)
    draw.pieslice((title_width - 110, 0, title_width + 110, 220), 180, 270, fill=GOLD)
    im.save(path)


def top_frame_asset(path: Path, width: int, height: int):
    im = Image.new("RGBA", (width, height), (0, 0, 0, 0))
    draw = ImageDraw.Draw(im)
    draw.rounded_rectangle((0, 70, width - 1, height - 1), radius=18, fill=CREAM, outline=(245, 228, 196), width=2)
    draw.rounded_rectangle((0, 0, 430, 100), radius=30, fill=GOLD)
    draw.rectangle((0, 62, 420, 120), fill=GOLD)
    draw.rounded_rectangle((18, 92, width - 18, height - 26), radius=4, fill=(255, 253, 248))
    im.save(path)


def save_hero(path: Path):
    base = Image.open(OLD_ORIGINAL).convert("RGBA")
    crop = base.crop((120, 685, 1005, 1210)).resize((884, 500))
    mask = Image.new("L", crop.size, 0)
    ImageDraw.Draw(mask).rounded_rectangle((0, 0, crop.size[0] - 1, crop.size[1] - 1), radius=22, fill=255)
    out = Image.new("RGBA", crop.size, (0, 0, 0, 0))
    out.paste(crop, (0, 0), mask)
    out.save(path)


def fmt_date_short(s: str) -> str:
    if re.match(r"^\d{4}[-/]\d{2}[-/]\d{2}$", s):
        d = datetime.strptime(s.replace("/", "-"), "%Y-%m-%d")
        return f"{d.year}/{d.month}/{d.day}"
    return s


def draw_center(draw: ImageDraw.ImageDraw, box, text: str, ft, fill=BLACK):
    x1, y1, x2, y2 = box
    bbox = draw.textbbox((0, 0), text, font=ft)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2 - 2), text, font=ft, fill=fill)


def save_table(path: Path, rows):
    w, h = 900, 190
    im = Image.new("RGBA", (w, h), (255, 252, 248, 255))
    draw = ImageDraw.Draw(im)
    cols = len(rows[0])
    col_widths = [210] + [(w - 210) // (cols - 1)] * (cols - 1)
    col_widths[-1] = w - sum(col_widths[:-1])
    row_h = h // len(rows)
    x = 0
    for c in range(cols):
        y = 0
        for r in range(len(rows)):
            box = (x, y, x + col_widths[c], y + row_h)
            fill = (244, 241, 235) if r == 0 else (255, 252, 248)
            draw.rectangle(box, fill=fill, outline=(64, 64, 64), width=2)
            txt = rows[r][c]
            if c == 0 and r in (1, 2):
                txt = fmt_date_short(txt)
            color = BLACK
            if r == 3 and c > 0:
                color = GREEN if float(txt) < 0 else RED
            draw_center(draw, box, txt, font(23), color)
            y += row_h
        x += col_widths[c]
    im.save(path)


def save_line_chart(path: Path, series, size, y_min, y_max, ticks, colors, legend_cols=3):
    w, h = size
    im = Image.new("RGBA", size, (0, 0, 0, 0))
    draw = ImageDraw.Draw(im)
    left, top, right, bottom = 70, 24, w - 35, h - 86
    draw.line((left, top, left, bottom), fill=(190, 190, 190), width=2)
    draw.line((left, bottom, right, bottom), fill=(190, 190, 190), width=2)
    for tick in ticks:
        yy = bottom - (tick - y_min) / (y_max - y_min) * (bottom - top)
        draw.line((left, yy, right, yy), fill=GRID, width=1)
        draw.text((6, yy - 11), f"{tick:.2f}", font=font(17), fill=GREY)
    all_dates = sorted({d for _, pairs in series for d, _ in pairs})
    start, end = all_dates[0], all_dates[-1]
    span = max((end - start).days, 1)
    for idx in range(0, len(all_dates), max(1, len(all_dates) // 8)):
        d = all_dates[idx]
        xx = left + (d - start).days / span * (right - left)
        draw.line((xx, bottom, xx, bottom + 8), fill=(205, 205, 205), width=2)
    for idx, (name, pairs) in enumerate(series):
        pts = []
        for d, val in pairs:
            xx = left + (d - start).days / span * (right - left)
            yy = bottom - (val - y_min) / (y_max - y_min) * (bottom - top)
            pts.append((xx, yy))
        if len(pts) > 1:
            draw.line(pts, fill=colors[idx % len(colors)], width=4, joint="curve")
    legend_y = h - 56
    col_w = w // legend_cols
    for idx, (name, _) in enumerate(series):
        lx = 45 + (idx % legend_cols) * col_w
        ly = legend_y + (idx // legend_cols) * 28
        label = name.replace("中债国债到期收益率:", "中债国债到期收益率:")
        label = label.replace("存款类机构质押式回购加权利率:", "回购加权利率:")
        draw.line((lx, ly + 10, lx + 44, ly + 10), fill=colors[idx % len(colors)], width=4)
        draw.text((lx + 50, ly), label, font=font(15), fill=GREY)
    im.save(path)


def text_after(prefix: str, paragraphs):
    for text in paragraphs:
        if text.startswith(prefix):
            return text
    return ""


def build_layout(content):
    assets = {}
    layers = []

    def asset_layer(section: str, name: str, path: Path, x: int, y: int):
        layers.append({"section": section, "type": "image", "name": name, "path": str(path), "x": x, "y": y})

    def add_text_layers(section: str, text_layers):
        for item in text_layers:
            item["section"] = section
            layers.append(item)

    # Body cover under new components. This is a plain background layer, not a final-image overlay.
    body_bg = ASSET_DIR / "body_background.png"
    bg_h = 6200
    Image.new("RGBA", (LONG_W, bg_h), DARK + (255,)).save(body_bg)
    asset_layer("00_bg", "主体深色背景", body_bg, 0, 500)

    date_patch = ASSET_DIR / "date_patch.png"
    Image.new("RGBA", (300, 92), (82, 68, 51, 255)).save(date_patch)
    asset_layer("00_header", "日期底板", date_patch, 420, 332)
    add_text_layers("00_header", [{"type": "text", "name": "日期_2026-5-21", "text": "2026-5-21", "x": 468, "y": 347, "size": 38, "color": (255, 221, 172)}])

    y = 505

    # Top summary.
    top_h = 1015
    top_frame = ASSET_DIR / "top_frame.png"
    top_frame_asset(top_frame, CARD_W, top_h)
    asset_layer("01_首屏摘要", "首屏卡片底板", top_frame, MARGIN, y)
    add_text_layers("01_首屏摘要", [{"type": "text", "name": "首屏标题", "text": "国债收益率整体下行", "x": 92, "y": y + 82, "size": 46, "color": BLACK}])
    asset_layer("01_首屏摘要", "首屏头图", ASSET_DIR / "hero.png", 120, y + 190)
    intro_segments = [
        ("上周债券收益率整体下行约2BP。资金利率小幅上行，央行公开市场净投放呵护月末资金面。", BLACK),
        ("期限上看，1年、5-30年国债下行更明显。", RED),
        ("债市多头趋势延续，但长端波动可能加大。建议控制风险，关注收益率阶段性上行后的配置机会。", BLACK),
    ]
    intro_layers, intro_h = rich_text_layers("首屏摘要", 90, y + 725, intro_segments, 34, 945, 56)
    add_text_layers("01_首屏摘要", intro_layers)
    y += top_h + 34

    # Bond performance.
    perf_h = 1240
    perf_frame = ASSET_DIR / "section_perf_frame.png"
    frame_asset(perf_frame, CARD_W, perf_h, 430)
    asset_layer("02_债市表现", "债市表现底板", perf_frame, MARGIN, y)
    add_text_layers("02_债市表现", [{"type": "text", "name": "债市表现标题", "text": "债市表现", "x": 130, "y": y + 17, "size": 50, "color": WHITE}])
    add_text_layers(
        "02_债市表现",
        [
            {"type": "text", "name": "债市表现一句话_日期", "text": "5月15日-5月21日，", "x": 78, "y": y + 145, "size": 34, "color": BLACK},
            {"type": "text", "name": "债市表现一句话_结论", "text": "国债收益率全期限下行。", "x": 430, "y": y + 145, "size": 34, "color": RED},
        ],
    )
    inner = ASSET_DIR / "perf_inner_card.png"
    rounded_asset(inner, (960, 950), 12, (255, 252, 248, 255), GOLD, 2)
    asset_layer("02_债市表现", "债市表现图表卡", inner, 82, y + 255)
    add_text_layers("02_债市表现", [{"type": "text", "name": "国债表格标题", "text": "不同期限国债收益率（上两周数据对比）", "x": 248, "y": y + 296, "size": 30, "color": BLACK}])
    asset_layer("02_债市表现", "国债收益率表格", ASSET_DIR / "yield_table.png", 112, y + 360)
    add_text_layers("02_债市表现", [{"type": "text", "name": "国债曲线标题", "text": "国债收益率曲线（近一年）", "x": 370, "y": y + 608, "size": 30, "color": BLACK}])
    asset_layer("02_债市表现", "国债收益率曲线", ASSET_DIR / "yield_chart.png", 125, y + 660)
    y += perf_h + 34

    # Analysis.
    analysis_h = 1260
    analysis_frame = ASSET_DIR / "section_analysis_frame.png"
    frame_asset(analysis_frame, CARD_W, analysis_h, 430)
    asset_layer("03_债市分析", "债市分析底板", analysis_frame, MARGIN, y)
    add_text_layers("03_债市分析", [{"type": "text", "name": "债市分析标题", "text": "债市分析", "x": 130, "y": y + 17, "size": 50, "color": WHITE}])
    analysis_inner = ASSET_DIR / "analysis_inner_card.png"
    rounded_asset(analysis_inner, (960, 1055), 12, (255, 252, 248, 255), GOLD, 2)
    asset_layer("03_债市分析", "债市分析内容卡", analysis_inner, 82, y + 150)
    add_text_layers("03_债市分析", [{"type": "text", "name": "债市分析主标题", "text": "影响债市波动的两大因素", "x": 330, "y": y + 205, "size": 32, "color": BLACK}])
    add_text_layers("03_债市分析", [{"type": "text", "name": "资金面标签", "text": "资金面", "x": 130, "y": y + 315, "size": 30, "color": BLACK}])
    fund_segments = [
        ("1、央行在公开市场操作上", BLACK),
        ("净投放1490亿元", RED),
        ("。\n2、", BLACK),
        ("资金利率上行", RED),
        ("，DR007收1.3280%；1年期国股行存单发行利率上行至1.445%。", BLACK),
    ]
    add_text_layers("03_债市分析", rich_text_layers("资金面正文", 315, y + 300, fund_segments, 28, 640, 46)[0])
    add_text_layers("03_债市分析", [{"type": "text", "name": "资金图标题", "text": "DR001/DR007上周情况：资金利率上行", "x": 285, "y": y + 485, "size": 28, "color": RED}])
    asset_layer("03_债市分析", "资金利率曲线", ASSET_DIR / "fund_chart.png", 190, y + 535)
    add_text_layers("03_债市分析", [{"type": "text", "name": "风险偏好标签", "text": "市场风险偏好", "x": 98, "y": y + 1010, "size": 28, "color": BLACK}])
    risk_segments = [
        ("1、权益高位回落后企稳，", BLACK),
        ("主力资金仍有兑现意愿", RED),
        ("。\n2、油价回落，黄金反弹有限；风险偏好边际变化对长端定价仍有影响。", BLACK),
    ]
    add_text_layers("03_债市分析", rich_text_layers("风险偏好正文", 315, y + 970, risk_segments, 28, 640, 46)[0])
    y += analysis_h + 34

    # Outlook.
    outlook_h = 410
    outlook_frame = ASSET_DIR / "section_outlook_frame.png"
    frame_asset(outlook_frame, CARD_W, outlook_h, 470)
    asset_layer("04_后市展望", "后市展望底板", outlook_frame, MARGIN, y)
    add_text_layers("04_后市展望", [{"type": "text", "name": "后市展望标题", "text": "后市展望", "x": 130, "y": y + 17, "size": 50, "color": WHITE}])
    outlook_segments = [
        ("展望后市，当前银行间资金利率上行符合季节性规律，不改", BLACK),
        ("资金面宽松的基本格局", RED),
        ("。", BLACK),
        ("债市多头趋势延续", RED),
        ("，且仍有较多机构存在欠配压力。建议在控制风险的前提下，把握收益率可能出现的阶段性上行机会，", BLACK),
        ("逢高配置", RED),
        ("。", BLACK),
    ]
    add_text_layers("04_后市展望", rich_text_layers("后市展望正文", 86, y + 150, outlook_segments, 32, 940, 52)[0])
    y += outlook_h + 34

    # Strategy.
    strategy_h = 510
    strategy_frame = ASSET_DIR / "section_strategy_frame.png"
    frame_asset(strategy_frame, CARD_W, strategy_h, 470)
    asset_layer("05_配置策略", "配置策略底板", strategy_frame, MARGIN, y)
    add_text_layers("05_配置策略", [{"type": "text", "name": "配置策略标题", "text": "配置策略", "x": 130, "y": y + 17, "size": 50, "color": WHITE}])
    strategy_text = "境内债券市场：\n建议以底仓价值和标配为主。\n1.波动容忍度较低的客户，建议主要配置稳健低波、短债等产品。\n2.对于存在一定波动容忍度的客户，可尝试优选“固收+”产品进行适当搭配。"
    add_text_layers("05_配置策略", plain_text_layers("配置策略正文", 90, y + 155, strategy_text, 34, 930, 58, BLACK)[0])
    y += strategy_h + 32

    # Source and risk.
    source_risk = f"{content['source']}\n{content['risk_warning']}"
    risk_layers, risk_h = plain_text_layers("风险提示", 28, y, source_risk, 24, 1068, 36, (126, 126, 126))
    add_text_layers("06_风险提示", risk_layers)
    y += risk_h + 70

    return layers, y


def main():
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    paragraphs, doc = get_paragraphs()
    table_rows = parse_table(doc)
    yield_series = parse_chart("chart1.xml")
    fund_series = parse_chart("chart2.xml")

    save_hero(ASSET_DIR / "hero.png")
    save_table(ASSET_DIR / "yield_table.png", table_rows)
    save_line_chart(
        ASSET_DIR / "yield_chart.png",
        yield_series,
        (880, 460),
        0.80,
        2.80,
        [0.80, 1.30, 1.80, 2.30, 2.80],
        [(86, 160, 220), (232, 122, 55), (165, 169, 170), (245, 180, 0), (68, 112, 196), (105, 168, 80)],
        legend_cols=3,
    )
    save_line_chart(
        ASSET_DIR / "fund_chart.png",
        fund_series,
        (760, 400),
        1.20,
        1.40,
        [1.20, 1.24, 1.28, 1.32, 1.36, 1.40],
        [(86, 160, 220), (239, 125, 45)],
        legend_cols=1,
    )

    content = {
        "docx": str(DOCX),
        "original_psd": str(PSD),
        "output_psd": str(OUT_DIR / "金葵花债市周观察20260521_基础v0.psd"),
        "output_png": str(OUT_DIR / "金葵花债市周观察20260521_基础v0.png"),
        "title": paragraphs[0],
        "subtitle": paragraphs[1],
        "intro_source": paragraphs[4],
        "performance": text_after("5月15日至5月21日", paragraphs),
        "risk_preference": text_after("市场风险偏好", paragraphs),
        "funding": text_after("资金面", paragraphs),
        "outlook": paragraphs[5],
        "source": text_after("数据来源", paragraphs),
        "risk_warning": text_after("风险提示", paragraphs),
        "x0": X0,
        "width": LONG_W,
    }
    layers, final_height = build_layout(content)
    content["layers"] = layers
    content["final_height"] = final_height

    (WORK_DIR / "content.json").write_text(json.dumps(content, ensure_ascii=False, indent=2), encoding="utf-8")
    print(WORK_DIR / "content.json")
    print(OUT_DIR)
    print(final_height)


if __name__ == "__main__":
    main()
