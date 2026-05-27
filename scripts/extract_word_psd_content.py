from __future__ import annotations

import json
import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document

from project_paths import ROOT, VARIANT_WORD_DIR, WORK_DIR

RULES_PATH = ROOT / "schemas" / "red_rules.json"
OUT_DIR = WORK_DIR / "psd_content"

VARIANT_DOCS = {
    "原版": VARIANT_WORD_DIR / "债市周观察原版.docx",
    "固收+": VARIANT_WORD_DIR / "债市周观察（固收+）.docx",
    "债市": VARIANT_WORD_DIR / "债市周观察（债市）.docx",
}


def nonempty_paragraphs(doc: Document) -> list[str]:
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def find_index(paragraphs: list[str], prefix: str) -> int:
    for idx, text in enumerate(paragraphs):
        if text.startswith(prefix):
            return idx
    raise ValueError(f"Missing paragraph prefix: {prefix}")


def find_text(paragraphs: list[str], prefix: str) -> str:
    return paragraphs[find_index(paragraphs, prefix)]


def next_after(paragraphs: list[str], prefix: str) -> str:
    idx = find_index(paragraphs, prefix)
    if idx + 1 >= len(paragraphs):
        raise ValueError(f"Missing paragraph after: {prefix}")
    return paragraphs[idx + 1]


def normalize_text(variant: str, field: str, text: str, rules: dict) -> tuple[str, list[dict]]:
    applied = []
    if field != "top_intro":
        return text, applied
    variant_rule = rules.get("normalizations", {}).get(variant)
    if not variant_rule:
        return text, applied
    old = variant_rule["from"]
    new = variant_rule["to"]
    if old in text:
        text = text.replace(old, new)
        applied.append({"from": old, "to": new})
    return text, applied


def split_rich_text(text: str, phrases: list[str]) -> list[dict]:
    phrases = sorted({p for p in phrases if p}, key=len, reverse=True)
    segments: list[dict] = []
    pos = 0

    while pos < len(text):
        match_pos = None
        match_phrase = None
        for phrase in phrases:
            idx = text.find(phrase, pos)
            if idx == -1:
                continue
            if match_pos is None or idx < match_pos or (idx == match_pos and len(phrase) > len(match_phrase or "")):
                match_pos = idx
                match_phrase = phrase

        if match_pos is None or match_phrase is None:
            segments.append({"text": text[pos:], "color": "black"})
            break

        if match_pos > pos:
            segments.append({"text": text[pos:match_pos], "color": "black"})

        segments.append({"text": match_phrase, "color": "red", "rule": match_phrase})
        pos = match_pos + len(match_phrase)

    return [seg for seg in segments if seg["text"]]


def table_rows(doc: Document) -> list[list[str]]:
    if not doc.tables:
        return []
    return [[cell.text.strip() for cell in row.cells] for row in doc.tables[0].rows]


def chart_xml_names(docx_path: Path) -> list[str]:
    with ZipFile(docx_path) as zf:
        names = zf.namelist()
    return sorted(
        name
        for name in names
        if name.startswith("word/charts/chart") and name.endswith(".xml")
    )


def extract_date(source: str) -> str:
    match = re.search(r"(\d{4})年(\d{1,2})月(\d{1,2})日", source)
    if not match:
        return ""
    year, month, day = match.groups()
    return f"{int(year)}-{int(month)}-{int(day)}"


def extract_week_range(sentence: str) -> str:
    match = re.search(r"(\d{1,2}月\d{1,2}日至\d{1,2}月\d{1,2}日)", sentence)
    return match.group(1) if match else ""


def strip_prefix(text: str, prefix: str) -> str:
    return text[len(prefix) :].strip() if text.startswith(prefix) else text


def top_title_from_subtitle(subtitle: str) -> str:
    subtitle = subtitle.lstrip("-—").strip()
    if subtitle.startswith("国债"):
        return subtitle
    return f"国债{subtitle}"


def build_variant(variant: str, docx_path: Path, rules: dict) -> dict:
    doc = Document(docx_path)
    paragraphs = nonempty_paragraphs(doc)

    title = paragraphs[0]
    subtitle = paragraphs[1].lstrip("-—").strip()

    first_section_idx = find_index(paragraphs, "第一板块")
    top_intro = paragraphs[first_section_idx + 1]
    top_intro, top_normalizations = normalize_text(variant, "top_intro", top_intro, rules)

    market_idx = find_index(paragraphs, "债市表现")
    market_sentence = paragraphs[market_idx + 1]
    yield_trend_title = paragraphs[market_idx + 2]
    yield_chart_caption = find_text(paragraphs, "图1.")
    yield_table_caption = find_text(paragraphs, "表1.")

    analysis_idx = find_index(paragraphs, "债市分析")
    risk_preference = paragraphs[analysis_idx + 1]
    funding = paragraphs[analysis_idx + 2]
    fund_chart_raw = find_text(paragraphs, "本周资金利率").replace("本周", "上周")
    fund_chart_state = fund_chart_raw.split("：", 1)[-1].strip()
    fund_chart_title = f"DR001/DR007上周情况：\n资金利率{fund_chart_state}"

    outlook_section_title = find_text(paragraphs, "二、后市展望")
    outlook = next_after(paragraphs, "二、后市展望")
    source = find_text(paragraphs, "数据来源")
    risk_warning = find_text(paragraphs, "风险提示")

    strategy = rules["fixed_strategy"]["text"]
    fields = {
        "title": title,
        "subtitle": subtitle,
        "top_title": top_title_from_subtitle(subtitle),
        "top_intro": top_intro,
        "market_section_title": "债市表现",
        "market_sentence": market_sentence,
        "yield_trend_title": yield_trend_title,
        "yield_chart_caption": yield_chart_caption,
        "yield_table_caption": yield_table_caption,
        "analysis_section_title": "债市分析",
        "risk_preference": strip_prefix(risk_preference, "市场风险偏好："),
        "funding": strip_prefix(funding, "资金面："),
        "fund_chart_title": fund_chart_title,
        "outlook_section_title": outlook_section_title.replace("二、", ""),
        "outlook": outlook,
        "strategy_section_title": "配置策略",
        "strategy": strategy,
        "source": source,
        "risk_warning": risk_warning,
    }

    field_rules = rules["field_rules"]
    rich_text = {
        field: split_rich_text(fields[field], field_rules.get(field, []))
        for field in (
            "top_intro",
            "market_sentence",
            "risk_preference",
            "funding",
            "fund_chart_title",
            "outlook",
            "strategy",
        )
    }

    charts = chart_xml_names(docx_path)
    return {
        "schema_version": "1.0",
        "meta": {
            "variant": variant,
            "source_docx": str(docx_path),
            "date": extract_date(source),
            "week_range": extract_week_range(market_sentence),
        },
        "layout": {
            "width": 1125,
            "adaptive_height": True,
            "height_policy": "measure_text_then_shift_following_sections",
        },
        "fields": fields,
        "rich_text": rich_text,
        "assets": {
            "yield_table": table_rows(doc),
            "yield_chart_xml": charts[0] if len(charts) > 0 else "",
            "fund_chart_xml": charts[1] if len(charts) > 1 else "",
        },
        "normalizations_applied": {
            "top_intro": top_normalizations,
        },
    }


def main() -> None:
    rules = json.loads(RULES_PATH.read_text(encoding="utf-8"))
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    index = []
    for variant, docx_path in VARIANT_DOCS.items():
        if not docx_path.exists():
            raise FileNotFoundError(docx_path)
        data = build_variant(variant, docx_path, rules)
        out_path = OUT_DIR / f"{variant}.json"
        out_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        index.append(
            {
                "variant": variant,
                "json": str(out_path),
                "source_docx": str(docx_path),
                "date": data["meta"]["date"],
                "week_range": data["meta"]["week_range"],
            }
        )
        print(out_path)

    index_path = OUT_DIR / "index.json"
    index_path.write_text(json.dumps(index, ensure_ascii=False, indent=2), encoding="utf-8")
    print(index_path)


if __name__ == "__main__":
    main()
