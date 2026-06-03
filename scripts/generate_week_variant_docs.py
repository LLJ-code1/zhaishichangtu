from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document

from project_paths import RAW_DOCX, VARIANT_WORD_DIR


VARIANT_FILENAMES = {
    "原版": "债市周观察原版.docx",
    "固收+": "债市周观察（固收+）.docx",
    "债市": "债市周观察（债市）.docx",
}

VARIANT_TOP_INTROS = {
    "固收+": (
        "本周债券收益率全期限下行2-4BP，5-10年期受到配置盘追捧，表现最优。"
        "央行净投放充裕，资金面维持整体宽松，叠加权益窄幅震荡、消费板块低位反弹，"
        "股债偏强共振下，固收+策略的配置价值有所提升。对于能接受一定净值波动的客户，"
        "可在底仓纯债基础上，关注固收+产品的阶段性配置机会，具体详情请见下文。"
    ),
    "债市": (
        "本周债券收益率全期限下行2-4BP，5-10年期受到配置盘追捧，表现最优。"
        "虽然中短端静态收益被大幅压缩，短期或仍有波动，但配置盘仍有欠配压力、"
        "交易盘积极转多，债市多头趋势或仍能延续。建议继续持有纯债产品，"
        "并在做好风险控制的前提下关注长端超额下行空间，具体详情请见下文。"
    ),
}


def nonempty_paragraph_indices(doc: Document) -> list[int]:
    return [idx for idx, paragraph in enumerate(doc.paragraphs) if paragraph.text.strip()]


def first_section_intro_index(doc: Document) -> int:
    indices = nonempty_paragraph_indices(doc)
    for pos, idx in enumerate(indices):
        if doc.paragraphs[idx].text.strip().startswith("第一板块"):
            if pos + 1 >= len(indices):
                raise ValueError("Missing top intro paragraph after 第一板块")
            return indices[pos + 1]
    raise ValueError("Missing 第一板块 paragraph")


def replace_paragraph_text(doc: Document, paragraph_index: int, text: str) -> None:
    paragraph = doc.paragraphs[paragraph_index]
    for run in paragraph.runs[1:]:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def build_variant_docs(source_docx: Path, output_dir: Path) -> dict[str, Path]:
    if not source_docx.exists():
        raise FileNotFoundError(source_docx)

    output_dir.mkdir(parents=True, exist_ok=True)
    outputs: dict[str, Path] = {}

    original = Document(source_docx)
    intro_index = first_section_intro_index(original)
    original_top_intro = original.paragraphs[intro_index].text.strip()

    for variant, filename in VARIANT_FILENAMES.items():
        out_path = output_dir / filename
        shutil.copy2(source_docx, out_path)

        if variant != "原版":
            doc = Document(out_path)
            replace_paragraph_text(doc, first_section_intro_index(doc), VARIANT_TOP_INTROS[variant])
            doc.save(out_path)
        else:
            doc = Document(out_path)
            replace_paragraph_text(doc, first_section_intro_index(doc), original_top_intro)
            doc.save(out_path)

        outputs[variant] = out_path

    return outputs


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate three Word variant docs for the current week.")
    parser.add_argument("--source-docx", type=Path, default=RAW_DOCX)
    parser.add_argument("--output-dir", type=Path, default=VARIANT_WORD_DIR)
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    outputs = build_variant_docs(args.source_docx, args.output_dir)
    for path in outputs.values():
        print(path)


if __name__ == "__main__":
    main()
